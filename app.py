#!/usr/bin/env python3
"""法务 Agent — 合同审核 Demo (D4)"""
from flask import Flask, request, jsonify, render_template_string
import os
import json
import sqlite3
import math
from dotenv import load_dotenv
from cherry_client import chat, chat_json, embed, test_connection
from checklists import CHECKLISTS, REVIEW_PROMPT
from feishu_client import (
    list_chats as feishu_list_chats,
    send_post as feishu_send_post,
    send_text as feishu_send_text,
    download_message_file as feishu_download_file,
    get_user_info as feishu_get_user_info,
)
from monitor import init_monitor

load_dotenv()

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB 上传上限

# === 配置 ===
CHERRYIN_API_KEY = os.environ.get("CHERRYIN_API_KEY", "")
CHERRYIN_BASE_URL = os.environ.get("CHERRYIN_BASE_URL", "https://express-ent-admin.cherryin.ai/v1")
LLM_MODEL = os.environ.get("LLM_MODEL", "agent/deepseek-v4-pro")
EMBED_MODEL = os.environ.get("EMBED_MODEL", "baai/bge-m3")

DB_PATH = os.path.join(os.environ.get("DATA_DIR", os.path.dirname(__file__)), "legal.db")
ALERT_CHAT_ID = os.environ.get("FEISHU_ALERT_CHAT_ID", "")  # 留空=跳过飞书推送

# === 合同类型 ===
CONTRACT_TYPES = ["采购合同", "销售合同-toB", "销售合同-toC", "人事合同"]

# === 分类 Prompt ===
CLASSIFY_PROMPT = """分析以下合同文本,判断属于哪一类:

可选类型 + 识别关键词:
1. 采购合同 — 采购、供应商、供货、采购方、商品、设备、货物、验收、交付
2. 销售合同-toB — SaaS、软件服务、企业版、授权、订阅、乙方提供服务、服务费
3. 销售合同-toC — 消费者、个人、用户、购买、商品、退换货、消费者权益
4. 人事合同 — 劳动合同、聘用、雇主、员工、工资、薪资、试用期、工作岗位、职责、社会保险、公积金、解除合同

判断规则(按优先级):
- 出现"劳动合同/聘用/员工/工资/薪资/试用期/工作岗位/社会保险"→ 人事合同
- 出现"采购/供应商/供货/货物/设备"→ 采购合同
- 出现"SaaS/软件服务/订阅/授权"→ 销售合同-toB
- 出现"消费者/个人购买/退换货"→ 销售合同-toC

合同文本(前 1500 字):
{text}

返回 JSON(不要其他文字):
{{
  "type": "采购合同|销售合同-toB|销售合同-toC|人事合同",
  "confidence": 0.0到1.0,
  "reason": "30字以内判断依据,引用关键关键词"
}}
"""

# === 数据库初始化 ===
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("""CREATE TABLE IF NOT EXISTS documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        doc_type TEXT,
        doc_name TEXT,
        content TEXT,
        chunk_count INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")
    c.execute("""CREATE TABLE IF NOT EXISTS chunks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        doc_id INTEGER,
        chunk_index INTEGER,
        chunk_text TEXT,
        embedding TEXT,
        FOREIGN KEY (doc_id) REFERENCES documents(id)
    )""")
    conn.commit()
    conn.close()

init_db()

# === 监控初始化 ===
init_monitor(app, service_name="legal-agent", db_path=DB_PATH, llm_test_fn=test_connection,
             alert_feishu_fn=feishu_send_post, alert_chat_id=ALERT_CHAT_ID)

# === 页面 ===
INDEX_HTML = """<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>法务 Agent · 企业运营智能化</title>
<style>
:root{
  --c-primary:#6366f1;--c-primary-2:#8b5cf6;--c-primary-3:#a855f7;
  --c-bg:#0f0f1a;--c-bg-2:#1a1a2e;--c-surface:rgba(255,255,255,.04);--c-surface-2:rgba(255,255,255,.08);
  --c-text:#e4e4e7;--c-text-dim:#a1a1aa;--c-text-muted:#71717a;
  --c-border:rgba(255,255,255,.08);--c-border-hover:rgba(139,92,246,.4);
  --c-green:#10b981;--c-amber:#f59e0b;--c-red:#ef4444;--c-blue:#3b82f6;
  --c-gold:#d4a843;
  --radius:16px;--radius-sm:10px;
  --shadow:0 8px 32px rgba(0,0,0,.3);--shadow-glow:0 0 40px rgba(139,92,246,.15);
}
*{margin:0;padding:0;box-sizing:border-box}
html{scroll-behavior:smooth}
body{
  font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Segoe UI","Kaiti SC","STKaiti","KaiTi","楷体",sans-serif;
  background:var(--c-bg);color:var(--c-text);line-height:1.6;overflow-x:hidden;
  min-height:100vh;
}
/* 装饰性背景 */
body::before{content:'';position:fixed;inset:0;z-index:-2;background:
  radial-gradient(ellipse 80% 50% at 20% 0%,rgba(99,102,241,.15),transparent),
  radial-gradient(ellipse 60% 50% at 80% 30%,rgba(168,85,247,.12),transparent),
  radial-gradient(ellipse 50% 50% at 50% 100%,rgba(139,92,246,.1),transparent),
  var(--c-bg)}
body::after{content:'';position:fixed;inset:0;z-index:-1;opacity:.4;
  background-image:linear-gradient(rgba(255,255,255,.015) 1px,transparent 1px),linear-gradient(90deg,rgba(255,255,255,.015) 1px,transparent 1px);
  background-size:60px 60px;mask-image:radial-gradient(ellipse 80% 60% at 50% 30%,#000,transparent)}

/* 顶部导航 */
.nav{position:sticky;top:0;z-index:100;backdrop-filter:blur(20px);background:rgba(15,15,26,.7);border-bottom:1px solid var(--c-border)}
.nav-inner{max-width:1100px;margin:0 auto;padding:16px 24px;display:flex;align-items:center;justify-content:space-between}
.nav-brand{font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Songti SC","STSong","SimSun","宋体",sans-serif;display:flex;align-items:center;gap:10px;font-size:17px;font-weight:700;color:var(--c-text);text-decoration:none}
.nav-brand-icon{width:36px;height:36px;border-radius:10px;background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3));display:flex;align-items:center;justify-content:center;font-size:18px;box-shadow:0 4px 12px rgba(99,102,241,.4)}
.nav-brand-text{background:linear-gradient(135deg,#fff,#a78bfa);-webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
.nav-links{display:flex;gap:4px;align-items:center}
.nav-links a{padding:8px 14px;border-radius:8px;font-size:13.5px;font-weight:500;color:var(--c-text-dim);text-decoration:none;transition:all .2s}
.nav-links a:hover{color:var(--c-text);background:var(--c-surface)}
.nav-status{display:flex;align-items:center;gap:6px;padding:6px 12px;border-radius:20px;background:rgba(16,185,129,.1);border:1px solid rgba(16,185,129,.2);font-size:12px;color:var(--c-green);font-weight:600}
.nav-status .dot{width:6px;height:6px;border-radius:50%;background:var(--c-green);box-shadow:0 0 8px var(--c-green);animation:pulse 2s infinite}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.5}}

/* 主容器 */
.wrap{max-width:1100px;margin:0 auto;padding:40px 24px 80px}

/* Hero */
.hero{position:relative;padding:60px 0 40px;text-align:center}
.hero-tag{display:inline-flex;align-items:center;gap:8px;padding:6px 14px;border-radius:20px;background:var(--c-surface-2);border:1px solid var(--c-border);font-size:12.5px;font-weight:600;color:var(--c-primary-3);margin-bottom:24px;letter-spacing:.5px}
.hero-tag::before{content:'';width:6px;height:6px;border-radius:50%;background:var(--c-primary-2);box-shadow:0 0 8px var(--c-primary-2)}
.hero h1{font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Songti SC","STSong","SimSun","宋体",sans-serif;font-size:clamp(38px,6vw,64px);font-weight:800;line-height:1.1;letter-spacing:-.02em;margin-bottom:20px}
.hero h1 .grad{background:linear-gradient(135deg,#818cf8 0%,#a78bfa 50%,#c084fc 100%);-webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
.hero p{font-size:17px;color:var(--c-text-dim);max-width:600px;margin:0 auto 36px;line-height:1.7}
.hero-cta{display:flex;gap:12px;justify-content:center;flex-wrap:wrap}
.btn{display:inline-flex;align-items:center;gap:8px;padding:13px 28px;border-radius:12px;font-size:14.5px;font-weight:600;text-decoration:none;transition:all .25s;cursor:pointer;border:none;font-family:inherit}
.btn-primary{background:linear-gradient(135deg,var(--c-primary),var(--c-primary-2));color:#fff;box-shadow:0 4px 20px rgba(99,102,241,.4)}
.btn-primary:hover{transform:translateY(-2px);box-shadow:0 8px 30px rgba(99,102,241,.5)}
.btn-ghost{background:var(--c-surface);color:var(--c-text);border:1px solid var(--c-border)}
.btn-ghost:hover{background:var(--c-surface-2);border-color:var(--c-border-hover)}

/* 统计卡片 */
.stats{display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:16px;margin:48px 0}
.stat{position:relative;padding:24px;border-radius:var(--radius);background:var(--c-surface);border:1px solid var(--c-border);overflow:hidden;transition:all .3s}
.stat::before{content:'';position:absolute;top:0;left:0;right:0;height:1px;background:linear-gradient(90deg,transparent,rgba(139,92,246,.4),transparent)}
.stat:hover{border-color:var(--c-border-hover);transform:translateY(-3px);box-shadow:var(--shadow-glow)}
.stat-label{font-size:12px;color:var(--c-text-muted);font-weight:600;letter-spacing:.5px;text-transform:uppercase;margin-bottom:8px}
.stat-value{font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue",sans-serif;font-size:30px;font-weight:800;letter-spacing:-.02em}
.stat-value .unit{font-size:14px;font-weight:500;color:var(--c-text-dim);margin-left:4px}
.stat-trend{font-size:12px;color:var(--c-green);margin-top:6px;font-weight:600}

/* 分区标题 */
.section{margin:56px 0 24px}
.section-head{display:flex;align-items:center;justify-content:space-between;margin-bottom:20px}
.section-title{font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Songti SC","STSong","SimSun","宋体",sans-serif;font-size:22px;font-weight:700;letter-spacing:-.01em;display:flex;align-items:center;gap:10px}
.section-title::before{content:'';width:4px;height:24px;border-radius:2px;background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3))}
.section-sub{font-size:13.5px;color:var(--c-text-muted)}

/* 功能卡片 */
.features{display:grid;grid-template-columns:repeat(auto-fit,minmax(300px,1fr));gap:16px}
.feat{position:relative;padding:28px;border-radius:var(--radius);background:var(--c-surface);border:1px solid var(--c-border);transition:all .3s;cursor:default;overflow:hidden}
.feat::after{content:'';position:absolute;top:-50%;right:-50%;width:200%;height:200%;background:radial-gradient(circle,rgba(139,92,246,.06),transparent 50%);opacity:0;transition:opacity .3s;pointer-events:none}
.feat:hover{border-color:var(--c-border-hover);transform:translateY(-4px)}
.feat:hover::after{opacity:1}
.feat-icon{width:48px;height:48px;border-radius:12px;display:flex;align-items:center;justify-content:center;font-size:22px;margin-bottom:16px;background:var(--c-surface-2)}
.feat-icon.teal{background:linear-gradient(135deg,rgba(99,102,241,.2),rgba(168,85,247,.2));box-shadow:0 0 20px rgba(99,102,241,.15)}
.feat-icon.green{background:linear-gradient(135deg,rgba(16,185,129,.2),rgba(52,211,153,.2));box-shadow:0 0 20px rgba(16,185,129,.15)}
.feat-icon.blue{background:linear-gradient(135deg,rgba(59,130,246,.2),rgba(96,165,250,.2));box-shadow:0 0 20px rgba(59,130,246,.15)}
.feat-icon.amber{background:linear-gradient(135deg,rgba(245,158,11,.2),rgba(251,191,36,.2));box-shadow:0 0 20px rgba(245,158,11,.15)}
.feat-icon.gold{background:linear-gradient(135deg,rgba(212,168,67,.2),rgba(251,191,36,.2));box-shadow:0 0 20px rgba(212,168,67,.15)}
.feat h3{font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Songti SC","STSong","SimSun","宋体",sans-serif;font-size:16.5px;font-weight:700;margin-bottom:8px}
.feat p{font-size:13.5px;color:var(--c-text-dim);line-height:1.65;margin-bottom:14px}
.feat-badge{display:inline-flex;align-items:center;gap:4px;padding:3px 10px;border-radius:12px;font-size:11.5px;font-weight:600}
.badge-done{background:rgba(16,185,129,.12);color:var(--c-green);border:1px solid rgba(16,185,129,.2)}

/* CTA 大卡片 */
.cta-card{margin-top:48px;padding:48px;border-radius:24px;background:linear-gradient(135deg,rgba(99,102,241,.1),rgba(168,85,247,.1));border:1px solid var(--c-border);text-align:center;position:relative;overflow:hidden}
.cta-card::before{content:'';position:absolute;inset:0;background:radial-gradient(circle at 50% 0%,rgba(139,92,246,.15),transparent 60%);pointer-events:none}
.cta-card h2{font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Songti SC","STSong","SimSun","宋体",sans-serif;font-size:26px;font-weight:700;margin-bottom:10px;position:relative}
.cta-card p{color:var(--c-text-dim);margin-bottom:24px;position:relative}
.cta-card .hero-cta{position:relative}

/* 技术栈 */
.tech{display:flex;gap:8px;flex-wrap:wrap;justify-content:center;margin-top:48px;padding-top:32px;border-top:1px solid var(--c-border)}
.tech-item{padding:6px 14px;border-radius:8px;background:var(--c-surface);border:1px solid var(--c-border);font-size:12px;color:var(--c-text-dim);font-weight:500;transition:all .2s}
.tech-item:hover{color:var(--c-text);border-color:var(--c-border-hover)}

/* 底部 */
.footer{text-align:center;padding:32px 0 0;color:var(--c-text-muted);font-size:12.5px}
.footer a{color:var(--c-text-dim);text-decoration:none}

/* 响应式 */
@media(max-width:640px){
  .nav-links{display:none}
  .hero{padding:40px 0 24px}
  .stats{grid-template-columns:repeat(2,1fr)}
  .features{grid-template-columns:1fr}
}
</style>
</head>
<body>
<nav class="nav">
  <div class="nav-inner">
    <a class="nav-brand" href="/">
      <span class="nav-brand-icon">⚖️</span>
      <span class="nav-brand-text">Legal Agent</span>
    </a>
    <div class="nav-links">
      <a href="/upload">合同审核</a>
      <a href="/icp">ICP 文档</a>
      <a href="/monitor">监控</a>
      <a href="/api/documents" target="_blank">知识库</a>
    </div>
    <div class="nav-status"><span class="dot"></span> 运行中</div>
  </div>
</nav>

<div class="wrap">
  <!-- Hero -->
  <section class="hero">
    <div class="hero-tag">D7 · 企业运营智能化 Demo</div>
    <h1>智能<span class="grad">合同审核</span><br>与法律风险把控</h1>
    <p>上传合同 PDF / 文本,AI 自动分类识别,79 项风险 Checklist 逐条审核,RAG 检索模板条款,输出修改建议与风险报告,支持 OCR 扫描件。</p>
    <div class="hero-cta">
      <a class="btn btn-primary" href="/upload">📄 上传合同</a>
      <a class="btn btn-ghost" href="/icp">📋 ICP 需求文档</a>
      <a class="btn btn-ghost" href="/api/documents" target="_blank">📚 知识库</a>
    </div>
  </section>

  <!-- 实时统计 -->
  <div class="stats" id="stats">
    <div class="stat">
      <div class="stat-label">合同模板</div>
      <div class="stat-value">4<span class="unit">类</span></div>
      <div class="stat-trend" style="color:var(--c-text-muted)">采购/toB/toC/人事</div>
    </div>
    <div class="stat">
      <div class="stat-label">风险检查点</div>
      <div class="stat-value">79<span class="unit">项</span></div>
      <div class="stat-trend" style="color:var(--c-text-muted)">逐条审核</div>
    </div>
    <div class="stat">
      <div class="stat-label">知识库文档</div>
      <div class="stat-value"><span id="s-docs">—</span><span class="unit">篇</span></div>
      <div class="stat-trend" id="s-chunks">加载中...</div>
    </div>
    <div class="stat">
      <div class="stat-label">服务状态</div>
      <div class="stat-value" style="font-size:18px;color:var(--c-green)">● Online</div>
      <div class="stat-trend" style="color:var(--c-text-muted)" id="s-uptime">端口 5003</div>
    </div>
  </div>

  <!-- 功能模块 -->
  <div class="section">
    <div class="section-head">
      <div class="section-title">功能模块</div>
      <div class="section-sub">7 天交付 · 49 个单元测试全绿</div>
    </div>
    <div class="features">
      <div class="feat">
        <div class="feat-icon teal">📄</div>
        <h3>合同分类</h3>
        <p>上传 PDF / TXT → AI 自动识别合同类型(采购 / toB / toC / 人事),返回类型 + 置信度 + 理由。</p>
        <span class="feat-badge badge-done">✅ 已上线</span>
      </div>
      <div class="feat">
        <div class="feat-icon amber">🔍</div>
        <h3>合同审核</h3>
        <p>79 项风险 Checklist 逐条审核,输出 pass / warn / fail 三级状态 + 修改建议 + 整体风险评级。</p>
        <span class="feat-badge badge-done">✅ 已上线</span>
      </div>
      <div class="feat">
        <div class="feat-icon green">📚</div>
        <h3>RAG 知识库</h3>
        <p>4 类合同模板入库,bge-m3 向量检索,审核时自动召回参考条款,支持相似度排序与引用回显。</p>
        <span class="feat-badge badge-done">✅ 已上线</span>
      </div>
      <div class="feat">
        <div class="feat-icon blue">🖨️</div>
        <h3>OCR 扫描件</h3>
        <p>pypdf 文本提取 + tesseract OCR 降级方案,扫描 PDF 自动识别文字(200 DPI, chi_sim + eng)。</p>
        <span class="feat-badge badge-done">✅ 已上线</span>
      </div>
      <div class="feat">
        <div class="feat-icon gold">📝</div>
        <h3>ICP 需求文档</h3>
        <p>AI 生成 ICP 备案外包需求文档(9 章结构),填表即生成,支持一键推送飞书给代理公司。</p>
        <span class="feat-badge badge-done">✅ 已上线</span>
      </div>
      <div class="feat">
        <div class="feat-icon teal">📤</div>
        <h3>飞书输出</h3>
        <p>审核结果 / ICP 文档一键推送飞书群,支持 Bot 交互式查询(审核 / 模板 / 帮助)。</p>
        <span class="feat-badge badge-done">✅ 已上线</span>
      </div>
    </div>
  </div>

  <!-- CTA -->
  <div class="cta-card">
    <h2>开始审核</h2>
    <p>上传你的第一份合同,体验 AI 分类 + 风险审核 + 模板引用全流程</p>
    <div class="hero-cta">
      <a class="btn btn-primary" href="/upload">📄 立即上传</a>
      <a class="btn btn-ghost" href="/icp">📋 生成 ICP 文档</a>
    </div>
  </div>

  <!-- 技术栈 -->
  <div class="tech">
    <span class="tech-item">Python Flask</span>
    <span class="tech-item">CherryIN API</span>
    <span class="tech-item">DeepSeek V4 Pro</span>
    <span class="tech-item">BGE-M3 Embedding</span>
    <span class="tech-item">SQLite + 向量检索</span>
    <span class="tech-item">Tesseract OCR</span>
    <span class="tech-item">飞书 OpenAPI</span>
  </div>

  <div class="footer">
    企业运营智能化 Demo · 法务 Agent v1.0 · 服务器 124.222.181.129:5003<br>
    基于 2026-07-09 线下拜访会议需求 · 7 天敏捷交付
  </div>
</div>

<script>
async function loadStats(){
  try{
    const [docsResp, statsResp] = await Promise.all([
      fetch('/api/documents').then(r=>r.json()).catch(()=>null),
      fetch('/api/stats').then(r=>r.json()).catch(()=>null)
    ]);
    if(docsResp && docsResp.documents){
      const totalChunks = docsResp.documents.reduce((s,d)=>s+(d.chunk_count||0),0);
      document.getElementById('s-docs').textContent = docsResp.documents.length;
      document.getElementById('s-chunks').textContent = totalChunks + ' chunks 已入库';
    }
    if(statsResp && statsResp.uptime_human){
      document.getElementById('s-uptime').textContent = '运行 ' + statsResp.uptime_human;
    }
  }catch(e){console.log('stats load failed',e)}
}
loadStats();
</script>
</body>
</html>"""

ICP_HTML = """<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>ICP 备案外包需求文档 · 法务 Agent</title>
<style>
:root{
  --c-primary:#6366f1;--c-primary-2:#8b5cf6;--c-primary-3:#a855f7;
  --c-bg:#0f0f1a;--c-bg-2:#1a1a2e;--c-surface:rgba(255,255,255,.04);--c-surface-2:rgba(255,255,255,.08);
  --c-text:#e4e4e7;--c-text-dim:#a1a1aa;--c-text-muted:#71717a;
  --c-border:rgba(255,255,255,.08);--c-border-hover:rgba(139,92,246,.4);
  --c-green:#10b981;--c-amber:#f59e0b;--c-red:#ef4444;--c-blue:#3b82f6;
  --radius:16px;--radius-sm:10px;
  --shadow:0 8px 32px rgba(0,0,0,.3);--shadow-glow:0 0 40px rgba(139,92,246,.15);
}
*{margin:0;padding:0;box-sizing:border-box}
html{scroll-behavior:smooth}
body{
  font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Segoe UI","Kaiti SC","STKaiti","KaiTi","楷体",sans-serif;
  background:var(--c-bg);color:var(--c-text);line-height:1.6;overflow-x:hidden;min-height:100vh;
}
body::before{content:'';position:fixed;inset:0;z-index:-2;background:
  radial-gradient(ellipse 80% 50% at 20% 0%,rgba(99,102,241,.15),transparent),
  radial-gradient(ellipse 60% 50% at 80% 30%,rgba(168,85,247,.12),transparent),
  radial-gradient(ellipse 50% 50% at 50% 100%,rgba(139,92,246,.1),transparent),
  var(--c-bg)}
.nav{position:sticky;top:0;z-index:100;backdrop-filter:blur(20px);background:rgba(15,15,26,.7);border-bottom:1px solid var(--c-border)}
.nav-inner{max-width:1100px;margin:0 auto;padding:16px 24px;display:flex;align-items:center;justify-content:space-between}
.nav-brand{display:flex;align-items:center;gap:10px;font-size:17px;font-weight:700;color:var(--c-text);text-decoration:none}
.nav-brand-icon{width:36px;height:36px;border-radius:10px;background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3));display:flex;align-items:center;justify-content:center;font-size:18px;box-shadow:0 4px 12px rgba(99,102,241,.4)}
.nav-brand-text{background:linear-gradient(135deg,#fff,#a78bfa);-webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
.nav-links{display:flex;gap:4px;align-items:center}
.nav-links a{padding:8px 14px;border-radius:8px;font-size:13.5px;font-weight:500;color:var(--c-text-dim);text-decoration:none;transition:all .2s}
.nav-links a:hover{color:var(--c-text);background:var(--c-surface)}
.nav-links a.active{color:var(--c-text);background:var(--c-surface-2)}
.wrap{max-width:1100px;margin:0 auto;padding:40px 24px 80px}
.card{background:var(--c-surface);border:1px solid var(--c-border);border-radius:var(--radius);padding:24px;margin-bottom:20px;backdrop-filter:blur(10px);position:relative;overflow:hidden}
.card::before{content:'';position:absolute;top:0;left:0;right:0;height:1px;background:linear-gradient(90deg,transparent,rgba(139,92,246,.4),transparent)}
.card h2{font-size:18px;margin-bottom:16px;color:var(--c-text);font-weight:700}
.muted{color:var(--c-text-muted);font-size:13px}

.form-group{margin-bottom:14px}
.form-group label{display:block;font-size:13px;font-weight:600;color:var(--c-text-dim);margin-bottom:6px}
.form-group label .req{color:var(--c-red)}
.form-group input,.form-group select,.form-group textarea{width:100%;padding:10px 14px;border:1px solid var(--c-border);border-radius:10px;font-size:14px;font-family:inherit;background:var(--c-surface-2);color:var(--c-text);outline:none;transition:border-color .2s}
.form-group input:focus,.form-group select:focus,.form-group textarea:focus{border-color:var(--c-border-hover)}
.form-group textarea{min-height:70px;resize:vertical}
.form-row{display:flex;gap:12px}
.form-row .form-group{flex:1}
.checkbox-group{display:flex;flex-wrap:wrap;gap:8px}
.checkbox-item{display:inline-flex;align-items:center;gap:6px;padding:8px 14px;background:var(--c-surface-2);border:1px solid var(--c-border);border-radius:10px;font-size:13px;cursor:pointer;color:var(--c-text);transition:all .2s}
.checkbox-item:hover{border-color:var(--c-border-hover)}
.checkbox-item input{width:auto;accent-color:var(--c-primary-2)}
.btn{padding:11px 26px;border:none;border-radius:10px;font-size:14px;font-weight:600;cursor:pointer;font-family:inherit;transition:all .25s}
.btn-primary{background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3));color:#fff;box-shadow:0 4px 16px rgba(99,102,241,.3)}
.btn-primary:hover:not(:disabled){transform:translateY(-2px);box-shadow:0 6px 24px rgba(99,102,241,.5)}
.btn-feishu{background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3));color:#fff;box-shadow:0 4px 16px rgba(99,102,241,.3)}
.btn-feishu:hover:not(:disabled){transform:translateY(-2px);box-shadow:0 6px 24px rgba(99,102,241,.5)}
.btn:disabled{opacity:.5;cursor:not-allowed;transform:none;box-shadow:none}
.result{margin-top:16px;padding:20px;background:var(--c-surface-2);border-radius:12px;border:1px solid var(--c-border);display:none}
.result h3{margin-bottom:12px;color:var(--c-text);font-size:15px;font-weight:700}
.result pre{white-space:pre-wrap;word-wrap:break-word;font-size:13.5px;line-height:1.8;max-height:600px;overflow-y:auto;color:var(--c-text);font-family:inherit}
.loading{text-align:center;padding:40px;color:var(--c-text-muted);display:none}
.loading .spin{display:inline-block;width:32px;height:32px;border:3px solid var(--c-surface-2);border-top:3px solid var(--c-primary-2);border-radius:50%;animation:spin 1s linear infinite}
@keyframes spin{to{transform:rotate(360deg)}}
.feishu-row{margin-top:16px;padding-top:16px;border-top:1px solid var(--c-border);display:flex;gap:8px;align-items:center;flex-wrap:wrap}
.feishu-row select{padding:10px 14px;border:1px solid var(--c-border);border-radius:10px;font-size:14px;min-width:220px;background:var(--c-surface-2);color:var(--c-text);font-family:inherit;outline:none}
.feishu-row select:focus{border-color:var(--c-border-hover)}
.result-msg{margin-top:8px;padding:10px;border-radius:8px;font-size:13px;display:none}
.result-msg.success{background:rgba(16,185,129,.1);border:1px solid rgba(16,185,129,.3);color:var(--c-green);display:block}
.result-msg.error{background:rgba(239,68,68,.1);border:1px solid rgba(239,68,68,.3);color:var(--c-red);display:block}
</style>
</head>
<body>
<div class="nav"><div class="nav-inner">
    <a class="nav-brand" href="/"><span class="nav-brand-icon">⚖️</span><span class="nav-brand-text">法务 Agent</span></a>
    <div class="nav-links">
        <a href="/">首页</a>
        <a href="/upload">合同审核</a>
        <a href="/icp" class="active">ICP 文档</a>
        <a href="/monitor">监控</a>
    </div>
</div></div>
<div class="wrap">
    <div class="card">
        <h2>📋 ICP 备案外包需求文档生成器</h2>
        <p class="muted" style="margin-bottom:16px">填写企业信息,AI 自动生成 ICP 备案外包需求文档</p>
        <form id="icp-form" onsubmit="return false">
            <div class="form-row">
                <div class="form-group">
                    <label>企业名称 <span class="req">*</span></label>
                    <input type="text" id="company_name" placeholder="北京科技有限公司">
                </div>
                <div class="form-group">
                    <label>统一社会信用代码</label>
                    <input type="text" id="credit_code" placeholder="91110000XXXXXXX">
                </div>
            </div>
            <div class="form-row">
                <div class="form-group">
                    <label>联系人</label>
                    <input type="text" id="contact_person" placeholder="张三">
                </div>
                <div class="form-group">
                    <label>联系电话</label>
                    <input type="text" id="contact_phone" placeholder="13800000000">
                </div>
            </div>
            <div class="form-row">
                <div class="form-group">
                    <label>网站名称 <span class="req">*</span></label>
                    <input type="text" id="site_name" placeholder="企业官网">
                </div>
                <div class="form-group">
                    <label>域名 <span class="req">*</span></label>
                    <input type="text" id="domain" placeholder="example.com">
                </div>
            </div>
            <div class="form-row">
                <div class="form-group">
                    <label>网站类型</label>
                    <select id="site_type">
                        <option>企业官网</option>
                        <option>电商平台</option>
                        <option>信息门户</option>
                        <option>SAAS 应用</option>
                        <option>行业论坛</option>
                    </select>
                </div>
                <div class="form-group">
                    <label>预计日均访问量</label>
                    <select id="daily_visits">
                        <option>1000 以下</option>
                        <option>1000-10000</option>
                        <option>10000-100000</option>
                        <option>10 万以上</option>
                    </select>
                </div>
            </div>
            <div class="form-group">
                <label>功能模块(多选)</label>
                <div class="checkbox-group">
                    <label class="checkbox-item"><input type="checkbox" value="信息发布" checked> 信息发布</label>
                    <label class="checkbox-item"><input type="checkbox" value="用户注册"> 用户注册</label>
                    <label class="checkbox-item"><input type="checkbox" value="在线支付"> 在线支付</label>
                    <label class="checkbox-item"><input type="checkbox" value="会员系统"> 会员系统</label>
                    <label class="checkbox-item"><input type="checkbox" value="搜索功能"> 搜索功能</label>
                    <label class="checkbox-item"><input type="checkbox" value="表单提交"> 表单提交</label>
                    <label class="checkbox-item"><input type="checkbox" value="文件下载"> 文件下载</label>
                </div>
            </div>
            <div class="form-group">
                <label>特殊要求</label>
                <textarea id="special_requirements" placeholder="如:需要 CDN 加速、多语言支持等"></textarea>
            </div>
            <button class="btn btn-primary" id="gen-btn" onclick="generateDoc()">🤖 生成需求文档</button>
        </form>
    </div>
    <div class="loading" id="loading">
        <div class="spin"></div>
        <p style="margin-top:10px">AI 正在生成需求文档,请稍候 10-30 秒...</p>
    </div>
    <div class="result" id="result">
        <h3>📄 ICP 备案外包需求文档</h3>
        <pre id="doc-content"></pre>
        <div class="feishu-row">
            <select id="chat-select"><option value="">加载群聊中...</option></select>
            <button class="btn btn-feishu" id="send-feishu" onclick="sendFeishu()" disabled>发送到飞书</button>
            <div class="result-msg" id="feishu-result"></div>
        </div>
    </div>
</div>
<script>
let lastDoc="";
function getFormData(){
    const features=[];
    document.querySelectorAll('.checkbox-item input:checked').forEach(c=>features.push(c.value));
    return{
        company_name:document.getElementById('company_name').value,
        credit_code:document.getElementById('credit_code').value,
        contact_person:document.getElementById('contact_person').value,
        contact_phone:document.getElementById('contact_phone').value,
        site_name:document.getElementById('site_name').value,
        domain:document.getElementById('domain').value,
        site_type:document.getElementById('site_type').value,
        daily_visits:document.getElementById('daily_visits').value,
        features:features,
        special_requirements:document.getElementById('special_requirements').value,
    };
}
async function generateDoc(){
    const data=getFormData();
    if(!data.company_name||!data.site_name||!data.domain){
        alert('请填写必填字段:企业名称、网站名称、域名');
        return;
    }
    document.getElementById('loading').style.display='block';
    document.getElementById('result').style.display='none';
    document.getElementById('gen-btn').disabled=true;
    try{
        const r=await fetch('/api/icp/generate',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(data)});
        const d=await r.json();
        if(d.ok){
            lastDoc=d.document;
            document.getElementById('doc-content').textContent=d.document;
            document.getElementById('result').style.display='block';
            loadChats();
        }else{
            alert('生成失败: '+d.error);
        }
    }catch(e){alert('网络错误');}
    document.getElementById('loading').style.display='none';
    document.getElementById('gen-btn').disabled=false;
}
async function loadChats(){
    try{
        const r=await fetch('/api/feishu/chats');
        const d=await r.json();
        const sel=document.getElementById('chat-select');
        if(d.ok&&d.chats&&d.chats.length>0){
            sel.innerHTML=d.chats.map(c=>'<option value="'+c.chat_id+'">'+c.name+'</option>').join('');
            document.getElementById('send-feishu').disabled=false;
        }else{sel.innerHTML='<option value="">无可用群聊</option>';}
    }catch(e){document.getElementById('chat-select').innerHTML='<option value="">加载失败</option>';}
}
async function sendFeishu(){
    const chatId=document.getElementById('chat-select').value;
    if(!chatId)return;
    const data=getFormData();
    const btn=document.getElementById('send-feishu');
    const result=document.getElementById('feishu-result');
    btn.disabled=true;btn.textContent='发送中...';
    result.className='result-msg';result.style.display='none';
    try{
        const r=await fetch('/api/icp/feishu',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({chat_id:chatId,form_data:data})});
        const d=await r.json();
        if(d.ok){result.className='result-msg success';result.textContent='✅ 已发送到飞书群';}
        else{result.className='result-msg error';result.textContent='❌ '+d.error;}
    }catch(e){result.className='result-msg error';result.textContent='❌ 网络错误';}
    btn.disabled=false;btn.textContent='发送到飞书';
}
</script>
</body>
</html>"""

UPLOAD_HTML = """<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>合同审核 · 法务 Agent</title>
<style>
:root{
  --c-primary:#6366f1;--c-primary-2:#8b5cf6;--c-primary-3:#a855f7;
  --c-bg:#0f0f1a;--c-bg-2:#1a1a2e;--c-surface:rgba(255,255,255,.04);--c-surface-2:rgba(255,255,255,.08);
  --c-text:#e4e4e7;--c-text-dim:#a1a1aa;--c-text-muted:#71717a;
  --c-border:rgba(255,255,255,.08);--c-border-hover:rgba(139,92,246,.4);
  --c-green:#10b981;--c-amber:#f59e0b;--c-red:#ef4444;--c-blue:#3b82f6;
  --radius:16px;--radius-sm:10px;
  --shadow:0 8px 32px rgba(0,0,0,.3);--shadow-glow:0 0 40px rgba(139,92,246,.15);
}
*{margin:0;padding:0;box-sizing:border-box}
html{scroll-behavior:smooth}
body{
  font-family:-apple-system,BlinkMacSystemFont,"Helvetica Neue","Segoe UI","Kaiti SC","STKaiti","KaiTi","楷体",sans-serif;
  background:var(--c-bg);color:var(--c-text);line-height:1.6;overflow-x:hidden;min-height:100vh;
}
body::before{content:'';position:fixed;inset:0;z-index:-2;background:
  radial-gradient(ellipse 80% 50% at 20% 0%,rgba(99,102,241,.15),transparent),
  radial-gradient(ellipse 60% 50% at 80% 30%,rgba(168,85,247,.12),transparent),
  radial-gradient(ellipse 50% 50% at 50% 100%,rgba(139,92,246,.1),transparent),
  var(--c-bg)}
.nav{position:sticky;top:0;z-index:100;backdrop-filter:blur(20px);background:rgba(15,15,26,.7);border-bottom:1px solid var(--c-border)}
.nav-inner{max-width:1100px;margin:0 auto;padding:16px 24px;display:flex;align-items:center;justify-content:space-between}
.nav-brand{display:flex;align-items:center;gap:10px;font-size:17px;font-weight:700;color:var(--c-text);text-decoration:none}
.nav-brand-icon{width:36px;height:36px;border-radius:10px;background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3));display:flex;align-items:center;justify-content:center;font-size:18px;box-shadow:0 4px 12px rgba(99,102,241,.4)}
.nav-brand-text{background:linear-gradient(135deg,#fff,#a78bfa);-webkit-background-clip:text;background-clip:text;-webkit-text-fill-color:transparent}
.nav-links{display:flex;gap:4px;align-items:center}
.nav-links a{padding:8px 14px;border-radius:8px;font-size:13.5px;font-weight:500;color:var(--c-text-dim);text-decoration:none;transition:all .2s}
.nav-links a:hover{color:var(--c-text);background:var(--c-surface)}
.nav-links a.active{color:var(--c-text);background:var(--c-surface-2)}
.wrap{max-width:1100px;margin:0 auto;padding:40px 24px 80px}
.card{background:var(--c-surface);border:1px solid var(--c-border);border-radius:var(--radius);padding:24px;margin-bottom:20px;backdrop-filter:blur(10px);position:relative;overflow:hidden}
.card::before{content:'';position:absolute;top:0;left:0;right:0;height:1px;background:linear-gradient(90deg,transparent,rgba(139,92,246,.4),transparent)}
.card h2{font-size:18px;margin-bottom:16px;color:var(--c-text);font-weight:700}
.muted{color:var(--c-text-muted);font-size:13px}

.drop-zone{border:2px dashed var(--c-border);border-radius:var(--radius);padding:48px;text-align:center;color:var(--c-text-muted);cursor:pointer;transition:all .3s;background:var(--c-surface-2)}
.drop-zone:hover{border-color:var(--c-border-hover);background:var(--c-surface);transform:scale(1.01)}
.drop-zone.dragover{border-color:var(--c-primary-2);background:var(--c-surface)}
.drop-zone.has-file{border-color:var(--c-green);background:rgba(16,185,129,.08);color:var(--c-green)}
.drop-icon{font-size:48px;margin-bottom:8px}
.drop-text{font-size:16px;font-weight:600;margin-bottom:4px;color:var(--c-text)}
.drop-hint{font-size:13px;opacity:.7}
.info-box{background:linear-gradient(135deg,rgba(99,102,241,.08),rgba(168,85,247,.05));border-left:3px solid var(--c-primary-2);border-radius:var(--radius-sm);padding:14px 18px;margin-top:16px;font-size:13px;color:var(--c-text-dim);line-height:1.7}
.info-box b{color:var(--c-text)}
.btn{display:inline-flex;align-items:center;gap:8px;padding:14px 36px;background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3));color:#fff;border:none;border-radius:12px;font-size:15px;font-weight:600;cursor:pointer;transition:all .25s;box-shadow:0 4px 16px rgba(99,102,241,.3);font-family:inherit}
.btn:hover:not(:disabled){transform:translateY(-2px);box-shadow:0 6px 24px rgba(99,102,241,.5)}
.btn:disabled{opacity:.5;cursor:not-allowed;transform:none;box-shadow:none}
.result-box{margin-top:16px}
.result-summary{background:linear-gradient(135deg,rgba(99,102,241,.1),rgba(168,85,247,.05));border:1px solid rgba(139,92,246,.2);border-radius:var(--radius-sm);padding:18px;margin-bottom:12px}
.risk-badge{display:inline-block;padding:4px 14px;border-radius:16px;font-size:13px;font-weight:700}
.risk-high{background:rgba(239,68,68,.15);color:var(--c-red);border:1px solid rgba(239,68,68,.3)}
.risk-medium{background:rgba(245,158,11,.15);color:var(--c-amber);border:1px solid rgba(245,158,11,.3)}
.risk-low{background:rgba(16,185,129,.15);color:var(--c-green);border:1px solid rgba(16,185,129,.3)}
.stat-pills{display:flex;gap:8px;margin-top:10px;flex-wrap:wrap}
.stat-pill{padding:4px 12px;border-radius:12px;font-size:12px;font-weight:600}
.pill-pass{background:rgba(16,185,129,.15);color:var(--c-green);border:1px solid rgba(16,185,129,.3)}
.pill-warn{background:rgba(245,158,11,.15);color:var(--c-amber);border:1px solid rgba(245,158,11,.3)}
.pill-fail{background:rgba(239,68,68,.15);color:var(--c-red);border:1px solid rgba(239,68,68,.3)}
.pill-total{background:rgba(99,102,241,.15);color:var(--c-primary-2);border:1px solid rgba(99,102,241,.3)}
.result-table{width:100%;border-collapse:collapse;font-size:13.5px;margin-top:12px}
.result-table th{background:var(--c-surface-2);color:var(--c-text-dim);padding:12px;text-align:left;font-weight:600;border-bottom:1px solid var(--c-border);text-transform:uppercase;letter-spacing:.3px;font-size:12.5px}
.result-table th:first-child{border-radius:8px 0 0 0}
.result-table th:last-child{border-radius:0 8px 0 0}
.result-table td{padding:12px;border-bottom:1px solid var(--c-border);vertical-align:top;color:var(--c-text)}
.result-table tr:hover td{background:var(--c-surface)}
.status-icon{font-size:18px}
.suggestion{color:var(--c-text-muted);font-size:12px;margin-top:4px}
.suggestion b{color:var(--c-text-dim)}
.loading{display:inline-block;width:20px;height:20px;border:3px solid var(--c-surface-2);border-top-color:var(--c-primary-2);border-radius:50%;animation:spin 1s linear infinite;margin-right:8px;vertical-align:middle}
@keyframes spin{to{transform:rotate(360deg)}}
.feishu-section{margin-top:20px;padding:20px;background:var(--c-surface-2);border-radius:var(--radius-sm);border:1px solid var(--c-border)}
.feishu-section h3{font-size:15px;margin-bottom:10px;color:var(--c-text);font-weight:700}
.feishu-section .chat-select{padding:10px 14px;border:1px solid var(--c-border);border-radius:10px;font-size:14px;margin-right:8px;min-width:220px;background:var(--c-surface);color:var(--c-text);font-family:inherit;outline:none}
.feishu-section .chat-select:focus{border-color:var(--c-border-hover)}
.btn-feishu-inline{background:linear-gradient(135deg,var(--c-primary),var(--c-primary-3));color:#fff;border:none;padding:10px 24px;border-radius:10px;font-size:14px;font-weight:600;cursor:pointer;font-family:inherit;box-shadow:0 4px 16px rgba(99,102,241,.3);transition:all .25s}
.btn-feishu-inline:hover:not(:disabled){transform:translateY(-2px);box-shadow:0 6px 24px rgba(99,102,241,.5)}
.btn-feishu-inline:disabled{opacity:.5;cursor:not-allowed;transform:none;box-shadow:none}
.result-msg{margin-top:10px;padding:10px;border-radius:8px;font-size:13px;display:none}
.result-msg.success{background:rgba(16,185,129,.1);border:1px solid rgba(16,185,129,.3);color:var(--c-green);display:block}
.result-msg.error{background:rgba(239,68,68,.1);border:1px solid rgba(239,68,68,.3);color:var(--c-red);display:block}
.error-box{color:var(--c-red);padding:16px;background:rgba(239,68,68,.08);border:1px solid rgba(239,68,68,.2);border-radius:10px}
.progress-box{text-align:center;padding:24px}
.progress-msg{margin-top:12px;color:var(--c-text-dim)}
.progress-sub{margin-top:8px;color:var(--c-text-muted);font-size:12px}
</style>
</head>
<body>
<div class="nav"><div class="nav-inner">
    <a class="nav-brand" href="/"><span class="nav-brand-icon">⚖️</span><span class="nav-brand-text">法务 Agent</span></a>
    <div class="nav-links">
        <a href="/">首页</a>
        <a href="/upload" class="active">合同审核</a>
        <a href="/icp">ICP 文档</a>
        <a href="/monitor">监控</a>
    </div>
</div></div>
<div class="wrap">
  <div class="card">
    <h2>📤 上传合同</h2>
    <p class="muted" style="margin-bottom:16px">上传合同 PDF/TXT → AI 分类 → 风险 Checklist 审核 → 结构化结论</p>
    <div class="drop-zone" id="drop">
      <div class="drop-icon">📁</div>
      <div class="drop-text">点击或拖拽文件到此处</div>
      <div class="drop-hint">支持 PDF 和 TXT 格式</div>
    </div>
    <input type="file" id="file" accept=".pdf,.txt" style="display:none">
    <div class="info-box">
      <b>📋 支持 4 类合同:</b> 采购合同 / 销售合同-toB(SaaS) / 销售合同-toC / 人事合同<br>
      <b>🔍 审核流程:</b> 分类 → 选 Checklist → RAG 检索模板 → 逐条审核 → 风险等级
    </div>
    <div style="text-align:center;margin-top:20px">
      <button class="btn" id="submit" disabled>🔍 分类 + 审核</button>
    </div>
    <div class="result-box" id="result"></div>
    <div class="feishu-section" id="feishu-section" style="display:none">
      <h3>📤 发送审核报告到飞书</h3>
      <p class="muted" style="margin-bottom:12px">选择群聊后,把审核结果发到飞书群</p>
      <select id="chat-select" class="chat-select"><option value="">加载群聊中...</option></select>
      <button id="send-feishu" class="btn-feishu-inline" disabled>发送到飞书</button>
      <div id="feishu-result" class="result-msg"></div>
    </div>
  </div>
</div>
<script>
function escapeHtml(s){if(s==null)return '';return String(s).replace(/[&<>"']/g,ch=>({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[ch]))}
const zone=document.getElementById('drop');
const file=document.getElementById('file');
const btn=document.getElementById('submit');
const result=document.getElementById('result');
zone.addEventListener('click',()=>file.click());
zone.addEventListener('dragover',e=>{e.preventDefault();zone.classList.add('dragover')});
zone.addEventListener('dragleave',e=>{zone.classList.remove('dragover')});
zone.addEventListener('drop',e=>{
  e.preventDefault();
  zone.classList.remove('dragover');
  if(e.dataTransfer.files.length){file.files=e.dataTransfer.files;showFile()}
});
file.addEventListener('change',showFile);
function showFile(){
  if(file.files.length){
    zone.innerHTML='<div class="drop-icon">✅</div><div class="drop-text">'+escapeHtml(file.files[0].name)+'</div><div class="drop-hint">点击重新选择</div>';
    zone.classList.add('has-file');
    btn.disabled=false;
  }
}
btn.addEventListener('click',async()=>{
  if(!file.files.length)return;
  result.innerHTML='<div class="progress-box"><span class="loading"></span><div class="progress-msg" id="progress-msg">⏳ 第 1 步:合同分类中(约 10-30 秒)...</div><div class="progress-sub">整体流程约 1-3 分钟,请耐心等待</div></div>';
  const fd=new FormData();
  fd.append('file',file.files[0]);
  const controller=new AbortController();
  const timeoutId=setTimeout(()=>controller.abort(),300000);
  const progressTimer=setTimeout(()=>{
    const pm=document.getElementById('progress-msg');
    if(pm)pm.innerHTML='⏳ 第 2 步:RAG 检索模板 + 逐条审核中(约 30-90 秒)...';
  },30000);
  try{
    const r=await fetch('/api/review/file',{method:'POST',body:fd,signal:controller.signal});
    clearTimeout(timeoutId);
    clearTimeout(progressTimer);
    const j=await r.json();
    if(j.items&&!j.error){
      const riskClass=j.overall_risk==='高'?'risk-high':j.overall_risk==='中'?'risk-medium':'risk-low';
      let html='<div class="result-summary">';
      html+='<div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px">';
      html+='<div><b>合同类型:</b> '+escapeHtml(j.contract_type)+' <span style="color:var(--c-text-muted);font-size:12px">('+(Number(j.type_confidence||0)*100).toFixed(0)+'% 置信度)</span></div>';
      html+='<span class="risk-badge '+riskClass+'">风险等级: '+escapeHtml(j.overall_risk)+'</span>';
      html+='</div>';
      html+='<div class="stat-pills">';
      html+='<span class="stat-pill pill-pass">✅ '+j.stats.pass+' 通过</span>';
      html+='<span class="stat-pill pill-warn">⚠️ '+j.stats.warn+' 警告</span>';
      html+='<span class="stat-pill pill-fail">❌ '+j.stats.fail+' 风险</span>';
      html+='<span class="stat-pill pill-total">共 '+j.stats.total+' 项</span>';
      html+='</div></div>';
      if(j.items&&j.items.length){
        html+='<table class="result-table"><thead><tr><th style="width:40px">状态</th><th>检查项</th><th>问题 / 建议</th></tr></thead><tbody>';
        j.items.forEach(it=>{
          const icon=it.status==='pass'?'✅':it.status==='warn'?'⚠️':'❌';
          const color=it.status==='pass'?'var(--c-green)':it.status==='warn'?'var(--c-amber)':'var(--c-red)';
          html+='<tr><td class="status-icon" style="color:'+color+'">'+icon+'</td><td>'+escapeHtml(it.item||'')+'</td><td>';
          if(it.issue)html+=escapeHtml(it.issue);
          if(it.suggestion)html+='<div class="suggestion"><b>建议:</b>'+escapeHtml(it.suggestion)+'</div>';
          html+='</td></tr>';
        });
        html+='</tbody></table>';
      }
      result.innerHTML=html;
      lastReview=j;
      const fs=document.getElementById('feishu-section');
      if(fs)fs.style.display='block';
    }else{
      result.innerHTML='<div class="error-box">❌ '+escapeHtml(j.error||JSON.stringify(j))+'</div>';
    }
  }catch(e){clearTimeout(timeoutId);clearTimeout(progressTimer);const msg=e.name==='AbortError'?'审核超时(超过 5 分钟),请重试或缩短合同文本':e.message;result.innerHTML='<div class="error-box">❌ 错误: '+escapeHtml(msg)+'</div>'}
});

let lastReview=null;
async function loadChats(){
  const sel=document.getElementById('chat-select');
  try{
    const r=await fetch('/api/feishu/chats');
    const d=await r.json();
    if(d.ok&&d.chats&&d.chats.length>0){
      sel.innerHTML=d.chats.map(c=>'<option value="'+escapeHtml(c.chat_id)+'">'+escapeHtml(c.name)+'</option>').join('');
      document.getElementById('send-feishu').disabled=false;
    }else{
      sel.innerHTML='<option value="">无可用群聊 ('+escapeHtml(d.error||'未知')+')</option>';
    }
  }catch(e){
    sel.innerHTML='<option value="">加载失败 ('+escapeHtml(e.message)+')</option>';
  }
}
document.getElementById('send-feishu').addEventListener('click',async()=>{
  if(!lastReview)return;
  const chatId=document.getElementById('chat-select').value;
  if(!chatId)return;
  const btn=document.getElementById('send-feishu');
  const msg=document.getElementById('feishu-result');
  btn.disabled=true;btn.textContent='发送中...';
  msg.className='result-msg';
  try{
    const r=await fetch('/api/review/feishu',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({chat_id:chatId,review:lastReview})});
    const d=await r.json();
    if(d.ok){msg.className='result-msg success';msg.textContent='✅ 审核报告已发送到飞书群';}
    else{msg.className='result-msg error';msg.textContent='❌ '+(d.error||'发送失败');}
  }catch(e){msg.className='result-msg error';msg.textContent='❌ '+e.message;}
  btn.disabled=false;btn.textContent='发送到飞书';
});
loadChats();
</script>
</body>
</html>"""


# === 路由 ===
@app.route("/")
def index():
    return INDEX_HTML

@app.route("/upload")
def upload():
    return UPLOAD_HTML

@app.route("/health")
def health():
    return jsonify({
        "status": "ok",
        "service": "legal-agent",
        "day": "D4",
        "port": 5003,
        "llm_model": LLM_MODEL,
        "embed_model": EMBED_MODEL,
    })

# === API ===
@app.route("/api/test-llm")
def test_llm():
    result = test_connection()
    return jsonify(result)

@app.route("/api/classify", methods=["POST"])
def classify():
    """合同类型分类(文本输入)"""
    data = request.get_json(silent=True) or {}
    text = data.get("text", "")
    if not text:
        return jsonify({"ok": False, "error": "text is required"})
    snippet = text[:1500]
    prompt = CLASSIFY_PROMPT.format(text=snippet)
    result = chat_json([
        {"role": "system", "content": "你是合同分类助手。只返回 JSON。"},
        {"role": "user", "content": prompt},
    ], temperature=0.1)
    if isinstance(result, dict) and result.get("_error"):
        return jsonify({"ok": False, "error": result["_error"], "stage": "classify"})
    if not isinstance(result, dict) or "type" not in result:
        return jsonify({"ok": False, "error": "LLM 返回格式异常", "stage": "classify", "raw": str(result)[:200]})
    return jsonify({"ok": True, "result": result})

@app.route("/api/classify-file", methods=["POST"])
def classify_file():
    """合同类型分类(文件上传)"""
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "file is required"})
    filename = f.filename
    text = extract_text_from_upload(f, filename)
    if isinstance(text, dict):  # 错误返回
        return jsonify(text)
    snippet = text[:1500]
    prompt = CLASSIFY_PROMPT.format(text=snippet)
    result = chat_json([
        {"role": "system", "content": "你是合同分类助手。只返回 JSON。"},
        {"role": "user", "content": prompt},
    ], temperature=0.1)
    if isinstance(result, dict) and result.get("_error"):
        return jsonify({"ok": False, "error": result["_error"], "stage": "classify"})
    if not isinstance(result, dict) or "type" not in result:
        return jsonify({"ok": False, "error": "LLM 返回格式异常", "stage": "classify", "raw": str(result)[:200]})
    return jsonify({"ok": True, "filename": filename, "text_length": len(text), "result": result})

# === RAG 知识库 ===
# OCR 配置
_OCR_MAX_PAGES = 50  # 单次 OCR 最多页数,防止超大 PDF 拖垮服务
_OCR_DPI = 200       # OCR 渲染 DPI(平衡速度与准确率)
_OCR_MIN_CHARS = 20  # pypdf 提取文本少于此字符数判定为扫描件,触发 OCR


def _ocr_pdf(pdf_bytes):
    """对扫描 PDF 做 OCR,返回提取的文本。

    用 pdf2image 把 PDF 每页转图片,再用 tesseract 识别。
    复用 OpenMentor 的 OCR 经验:200 DPI + chi_sim+eng。
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError as e:
        return None, f"OCR 依赖未安装: {e}"

    try:
        images = convert_from_bytes(pdf_bytes, dpi=_OCR_DPI, first_page=1,
                                    last_page=_OCR_MAX_PAGES)
    except Exception as e:
        return None, f"PDF 转图片失败: {e}"

    if not images:
        return None, "PDF 无可渲染页面"

    pages_text = []
    for i, img in enumerate(images, 1):
        try:
            page_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            pages_text.append(page_text.strip())
        except Exception as e:
            pages_text.append("")  # 单页失败不中断,继续后续页
        # 释放图片内存
        img.close()

    text = "\n\n".join(t for t in pages_text if t)
    return text, None


def extract_text_from_upload(f, filename):
    """从上传文件提取文本。成功返回 str,失败返回 dict(错误响应)。

    策略:
      1. TXT: 直接 UTF-8 解码
      2. PDF: 先 pypdf 提取文本;若文本过短(扫描件),自动降级 OCR
      3. 其他: 尝试 UTF-8 解码
    """
    raw_bytes = f.read()

    if filename.lower().endswith(".txt"):
        text = raw_bytes.decode("utf-8", errors="ignore")
        if not text.strip():
            return {"ok": False, "error": "文件为空"}
        return text

    if filename.lower().endswith(".pdf"):
        # 第一步:pypdf 提取文本
        text = ""
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(raw_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            return {"ok": False, "error": "pypdf not installed"}
        except Exception as e:
            return {"ok": False, "error": f"PDF 解析失败: {e}"}

        # 第二步:文本过短 → 判定为扫描件,触发 OCR
        if len(text.strip()) < _OCR_MIN_CHARS:
            ocr_text, ocr_err = _ocr_pdf(raw_bytes)
            if ocr_err:
                return {"ok": False, "error": f"扫描件需 OCR,但 OCR 失败: {ocr_err}"}
            if not ocr_text.strip():
                return {"ok": False, "error": "OCR 未提取到文本(图片可能模糊或为纯图形)"}
            return ocr_text

        return text

    if filename.lower().endswith(".docx"):
        # .docx 本质是 ZIP,直接 UTF-8 解码会得到乱码
        # 优先用 python-docx 提取;降级用 zipfile 直接读 word/document.xml
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(raw_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 也提取表格里的文字(合同经常用表格)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = chr(10).join(paragraphs)
            if not text.strip():
                return {"ok": False, "error": ".docx 文件为空或无文本内容"}
            return text
        except ImportError:
            # python-docx 未安装,降级用 zipfile 直接解析
            try:
                import zipfile, re
                with zipfile.ZipFile(io.BytesIO(raw_bytes)) as z:
                    xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
                    # 去掉 XML 标签,保留纯文本
                    text = re.sub(r"<w:p[ >]", chr(10), xml)
                    text = re.sub(r"<[^>]+>", "", text)
                    text = re.sub(r"\n{3,}", chr(10)+chr(10), text).strip()
                    if not text:
                        return {"ok": False, "error": ".docx 解析失败:文档无文本内容"}
                    return text
            except Exception as e:
                return {"ok": False, "error": f".docx 解析失败(需安装 python-docx 或检查文件): {e}"}
        except Exception as e:
            return {"ok": False, "error": f".docx 解析失败: {e}"}

    # 其他格式:尝试 UTF-8
    text = raw_bytes.decode("utf-8", errors="ignore")
    if not text.strip():
        return {"ok": False, "error": "无法提取文本(不支持的文件格式或文件为空)"}
    return text


@app.route("/api/documents")
def documents():
    """列出已入库文档"""
    conn = sqlite3.connect(DB_PATH)
    try:
        c = conn.cursor()
        c.execute("SELECT id, doc_type, doc_name, chunk_count, created_at FROM documents ORDER BY id DESC")
        rows = c.fetchall()
    finally:
        conn.close()
    return jsonify({"ok": True, "documents": [
        {"id": r[0], "doc_type": r[1], "doc_name": r[2], "chunk_count": r[3], "created_at": r[4]}
        for r in rows
    ]})

@app.route("/api/ingest", methods=["POST"])
def ingest():
    """文本入库(分块 + embedding),支持重复入库时先删旧数据"""
    data = request.get_json(silent=True) or {}
    doc_type = data.get("doc_type", "").strip()
    doc_name = data.get("doc_name", "").strip()
    content = data.get("content", "")
    if not doc_type:
        return jsonify({"ok": False, "error": "doc_type is required"})
    if doc_type not in CONTRACT_TYPES:
        return jsonify({"ok": False, "error": f"doc_type must be one of {CONTRACT_TYPES}"})
    if not doc_name:
        return jsonify({"ok": False, "error": "doc_name is required"})
    if not content or not content.strip():
        return jsonify({"ok": False, "error": "content is required"})
    chunk_size = 500
    overlap = 50
    if overlap >= chunk_size:
        return jsonify({"ok": False, "error": "overlap must be less than chunk_size"})
    chunks = []
    start = 0
    while start < len(content):
        end = start + chunk_size
        chunks.append(content[start:end])
        start = end - overlap
    # 去除尾部空 chunk(overlap 导致的)
    if len(chunks) > 1 and not chunks[-1].strip():
        chunks.pop()

    conn = sqlite3.connect(DB_PATH)
    try:
        c = conn.cursor()
        # 重复入库检测:同名文档先删旧数据
        c.execute("SELECT id FROM documents WHERE doc_name = ?", (doc_name,))
        old = c.fetchall()
        if old:
            for (old_id,) in old:
                c.execute("DELETE FROM chunks WHERE doc_id = ?", (old_id,))
                c.execute("DELETE FROM documents WHERE id = ?", (old_id,))

        c.execute("INSERT INTO documents (doc_type, doc_name, content, chunk_count) VALUES (?, ?, ?, ?)",
                  (doc_type, doc_name, content, len(chunks)))
        doc_id = c.lastrowid
        # embedding(批量,每批 10 个)
        for i in range(0, len(chunks), 10):
            batch = chunks[i:i+10]
            emb_result = embed(batch)
            if emb_result.get("error"):
                conn.rollback()
                return jsonify({"ok": False, "error": emb_result["error"], "stage": "embedding"})
            if not emb_result.get("embeddings"):
                conn.rollback()
                return jsonify({"ok": False, "error": "embedding API 返回空结果", "stage": "embedding"})
            for j, e in enumerate(emb_result["embeddings"]):
                c.execute("INSERT INTO chunks (doc_id, chunk_index, chunk_text, embedding) VALUES (?, ?, ?, ?)",
                          (doc_id, i+j, batch[j], json.dumps(e)))
        conn.commit()
    finally:
        conn.close()
    return jsonify({"ok": True, "doc_id": doc_id, "chunk_count": len(chunks)})

@app.route("/api/search", methods=["POST"])
def search():
    """向量检索"""
    data = request.get_json(silent=True) or {}
    query = data.get("query", "")
    if not query:
        return jsonify({"ok": False, "error": "query is required"})
    try:
        top_k = min(int(data.get("top_k", 3)), 20)  # 上限 20 防滥用
    except (ValueError, TypeError):
        top_k = 3
    emb_result = embed(query)
    if emb_result.get("error"):
        return jsonify({"ok": False, "error": emb_result["error"]})
    if not emb_result.get("embeddings"):
        return jsonify({"ok": False, "error": "embedding 返回空"})
    query_vec = emb_result["embeddings"][0]
    conn = sqlite3.connect(DB_PATH)
    try:
        c = conn.cursor()
        c.execute("""SELECT ch.id, ch.doc_id, ch.chunk_index, ch.chunk_text, ch.embedding,
                            d.doc_name, d.doc_type
                     FROM chunks ch JOIN documents d ON ch.doc_id = d.id""")
        scored = []
        norm_q = math.sqrt(sum(x * x for x in query_vec))
        for row in c.fetchall():
            emb = json.loads(row[4])
            if len(emb) != len(query_vec):
                continue  # 维度不一致跳过
            dot = sum(a * b for a, b in zip(query_vec, emb))
            norm_e = math.sqrt(sum(x * x for x in emb))
            sim = dot / (norm_q * norm_e) if norm_q > 0 and norm_e > 0 else 0
            scored.append((sim, row))
        scored.sort(key=lambda x: -x[0])
    finally:
        conn.close()
    results = []
    for sim, row in scored[:top_k]:
        text_preview = row[3][:200] + ("..." if len(row[3]) > 200 else "")
        results.append({
            "score": round(sim, 4),
            "chunk_index": row[2],
            "chunk_text": text_preview,
            "doc_name": row[5] or "",
            "doc_type": row[6] or "",
        })
    return jsonify({"ok": True, "query": query, "results": results})

def _do_review(text):
    """合同审核核心逻辑: 分类 → 选 Checklist → RAG(同类模板) → 逐条审核"""
    # 诊断日志:记录提取的文本长度和前 500 字
    import logging
    logging.getLogger("app").info(f"[_do_review] 提取文本长度={len(text)}, 前500字={text[:500]!r}")
    # 1. 分类
    snippet = text[:1500]
    classify_prompt = CLASSIFY_PROMPT.format(text=snippet)
    classify_result = chat_json([
        {"role": "system", "content": "你是合同分类助手。只返回 JSON。"},
        {"role": "user", "content": classify_prompt},
    ], temperature=0.1, timeout=180)

    # 分类失败显式返回错误,不静默回退
    if not isinstance(classify_result, dict) or classify_result.get("_error"):
        return {"ok": False, "error": f"合同分类失败: {classify_result.get('_error', 'invalid response') if isinstance(classify_result, dict) else 'non-dict'}", "stage": "classify"}

    contract_type = str(classify_result.get("type", "")).strip()
    confidence = classify_result.get("confidence", 0)
    logging.getLogger("app").info(f"[_do_review] 分类结果: type={contract_type!r} confidence={confidence} raw={classify_result}")
    # 数字化 confidence 并限制到 [0, 1] 范围
    try:
        confidence = float(confidence)
        confidence = max(0.0, min(1.0, confidence))
    except (ValueError, TypeError):
        confidence = 0

    if contract_type not in CHECKLISTS:
        return {"ok": False, "error": f"不支持的合同类型: {contract_type}", "stage": "classify", "raw_type": contract_type}

    # 2. 选 Checklist
    checklist_items = CHECKLISTS[contract_type]
    checklist_text = "\n".join(f"{i+1}. {item}" for i, item in enumerate(checklist_items))

    # 3. RAG 检索模板条款(只检索同类型模板,避免跨类型污染)
    rag_context = ""
    try:
        emb_result = embed(text[:500])
        if emb_result.get("error"):
            rag_context = f"(RAG 检索失败: {emb_result['error']})"
        elif emb_result.get("embeddings"):
            query_vec = emb_result["embeddings"][0]
            scored = []
            conn = sqlite3.connect(DB_PATH)
            try:
                c = conn.cursor()
                # 只检索同类型合同的模板
                c.execute("""
                    SELECT ch.chunk_text, ch.embedding
                    FROM chunks ch
                    JOIN documents d ON ch.doc_id = d.id
                    WHERE d.doc_type = ?
                """, (contract_type,))
                norm_q = math.sqrt(sum(x * x for x in query_vec))
                for row in c.fetchall():
                    emb = json.loads(row[1])
                    if len(emb) != len(query_vec):
                        continue  # 维度不一致跳过
                    dot = sum(a * b for a, b in zip(query_vec, emb))
                    norm_e = math.sqrt(sum(x * x for x in emb))
                    sim = dot / (norm_q * norm_e) if norm_q > 0 and norm_e > 0 else 0
                    scored.append((sim, row[0]))
            finally:
                conn.close()
            scored.sort(key=lambda x: -x[0])
            top_chunks = [s[1] for s in scored[:3]]
            if top_chunks:
                rag_context = "\n---\n".join(top_chunks)
            else:
                rag_context = f"(无 {contract_type} 类型模板)"
    except Exception as e:
        rag_context = f"(RAG 检索失败: {e})"

    # 4. 审核(送全文,deepseek-v4-pro 支持 128k 上下文)
    review_prompt = REVIEW_PROMPT.format(
        contract_type=contract_type,
        checklist=checklist_text,
        contract_text=text,
        rag_context=rag_context[:2000] or "(无)",
    )

    review_result = chat_json([
        {"role": "system", "content": "你是法务审核专员。逐条检查 Checklist,返回 JSON 数组。"},
        {"role": "user", "content": review_prompt},
    ], temperature=0.1, timeout=180)

    # 审核失败显式返回错误,不返回空列表 + 风险「低」
    if isinstance(review_result, dict) and review_result.get("_error"):
        return {"ok": False, "error": f"审核 LLM 调用失败: {review_result['_error']}", "stage": "review", "contract_type": contract_type}

    # 支持 LLM 返回数组 [...] 或 {"items": [...]}
    if isinstance(review_result, list):
        items = review_result
    elif isinstance(review_result, dict) and "items" in review_result:
        items = review_result["items"] if isinstance(review_result["items"], list) else []
    else:
        return {"ok": False, "error": "审核 LLM 返回格式异常,既非数组也非 {items:...}", "stage": "review", "contract_type": contract_type, "raw": str(review_result)[:200]}

    # 5. 统计(过滤非 dict 元素)
    valid_items = [it for it in items if isinstance(it, dict)]
    stats = {"pass": 0, "warn": 0, "fail": 0, "total": len(valid_items)}
    for item in valid_items:
        status = item.get("status", "warn")
        if status in ("pass", "warn", "fail"):
            stats[status] += 1
    overall_risk = "高" if stats["fail"] > 2 else ("中" if stats["fail"] > 0 or stats["warn"] > 3 else "低")

    return {
        "ok": True,
        "contract_type": contract_type,
        "type_confidence": confidence,
        "overall_risk": overall_risk,
        "stats": stats,
        "items": valid_items,
    }

@app.route("/api/review", methods=["POST"])
def review():
    """合同审核: 分类 → 选 Checklist → RAG 检索模板 → 逐条审核"""
    data = request.get_json(silent=True) or {}
    text = data.get("text", "")
    if not text:
        return jsonify({"ok": False, "error": "text is required"})
    return jsonify(_do_review(text))

@app.route("/api/review/file", methods=["POST"])
def review_file():
    """合同审核(文件上传)"""
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "error": "file is required"})
    filename = f.filename
    text = extract_text_from_upload(f, filename)
    if isinstance(text, dict):  # 错误返回
        return jsonify(text)
    return jsonify(_do_review(text))


# === D4: 飞书输出 ===
@app.route("/api/feishu/chats")
def feishu_chats():
    """列出 Bot 所在的飞书群聊"""
    return jsonify(feishu_list_chats())

@app.route("/api/review/feishu", methods=["POST"])
def review_feishu():
    """把审核结果发到飞书群"""
    data = request.get_json(silent=True) or {}
    chat_id = data.get("chat_id", "")
    review_data = data.get("review", {})
    if not chat_id:
        return jsonify({"ok": False, "error": "chat_id is required"})
    if not review_data:
        return jsonify({"ok": False, "error": "review data is required"})

    # 构建富文本卡片
    contract_type = review_data.get("contract_type", "未知")
    overall_risk = review_data.get("overall_risk", "未知")
    stats = review_data.get("stats", {})
    items = review_data.get("items", [])

    paragraphs = []
    # 概要行
    risk_emoji = {"高": "🔴", "中": "🟡", "低": "🟢"}.get(overall_risk, "⚪")
    paragraphs.append([{
        "tag": "text",
        "text": f"合同类型: {contract_type}\n总体风险: {risk_emoji} {overall_risk}\n统计: ✅{stats.get('pass',0)} 通过  ⚠️{stats.get('warn',0)} 警告  ❌{stats.get('fail',0)} 不合规\n",
    }])

    # 风险项详情(过滤非 dict 元素)
    fail_items = [it for it in items if isinstance(it, dict) and it.get("status") == "fail"]
    warn_items = [it for it in items if isinstance(it, dict) and it.get("status") == "warn"]
    if fail_items:
        seg = [{"tag": "text", "text": "\n❌ 不合规项:\n"}]
        for it in fail_items[:5]:
            seg.append({"tag": "text", "text": f"  • {str(it.get('item',''))[:40]}\n    建议: {str(it.get('suggestion',''))[:60]}\n"})
        paragraphs.append(seg)
    if warn_items:
        seg = [{"tag": "text", "text": "\n⚠️ 警告项:\n"}]
        for it in warn_items[:5]:
            seg.append({"tag": "text", "text": f"  • {str(it.get('item',''))[:40]}\n    建议: {str(it.get('suggestion',''))[:60]}\n"})
        paragraphs.append(seg)

    result = feishu_send_post(chat_id, "🔍 合同审核报告", paragraphs)
    if result.get("ok"):
        return jsonify({"ok": True, "sent": True})
    return jsonify({"ok": False, "error": result.get("error", "send failed"), "raw": result})


# === D6: ICP 外包需求文档 ===
ICP_PROMPT = """你是企业法务顾问。根据以下企业信息,生成一份 ICP 备案外包需求文档。

企业信息:
- 企业名称: {company_name}
- 统一社会信用代码: {credit_code}
- 联系人: {contact_person}
- 联系电话: {contact_phone}

网站信息:
- 网站名称: {site_name}
- 域名: {domain}
- 网站类型: {site_type}
- 预计日均访问量: {daily_visits}
- 功能模块: {features}
- 特殊要求: {special_requirements}

请按以下结构输出(使用 Markdown 格式):

# ICP 备案外包需求文档

## 一、项目概述
（简述企业背景和备案需求）

## 二、企业信息
（企业名称、信用代码、联系人等）

## 三、网站信息
（网站名称、域名、类型、访问量等）

## 四、功能需求
（根据选择的功能模块,逐项描述需求）

## 五、ICP 备案要求
（备案类型、所需材料、审批流程、时间预估）

## 六、技术要求
（服务器、带宽、安全等,基于访问量和功能）

## 七、服务范围
（外包方应提供的服务清单）

## 八、交付标准与验收
（交付物、验收标准）

## 九、时间安排
（备案周期预估,分阶段时间表）

要求:
1. 内容具体、可执行,不要泛泛而谈
2. 符合中国 ICP 备案相关法规要求
3. 根据企业实际情况定制(如电商需要特别注意什么)
"""


def _generate_icp_doc(form_data):
    """生成 ICP 外包需求文档,返回 (doc_text, error)"""
    prompt = ICP_PROMPT.format(
        company_name=form_data.get("company_name", ""),
        credit_code=form_data.get("credit_code", ""),
        contact_person=form_data.get("contact_person", ""),
        contact_phone=form_data.get("contact_phone", ""),
        site_name=form_data.get("site_name", ""),
        domain=form_data.get("domain", ""),
        site_type=form_data.get("site_type", ""),
        daily_visits=form_data.get("daily_visits", ""),
        features=", ".join(form_data.get("features", [])),
        special_requirements=form_data.get("special_requirements", "无"),
    )
    result = chat([
        {"role": "system", "content": "你是企业法务顾问,擅长 ICP 备案流程和需求文档撰写。"},
        {"role": "user", "content": prompt},
    ], temperature=0.3, timeout=120)
    if result.get("error"):
        return None, result["error"]
    doc = result.get("content", "").strip()
    if not doc:
        return None, "AI 返回空内容"
    return doc, None


@app.route("/icp")
def icp_page():
    """ICP 外包需求文档页面"""
    return ICP_HTML


@app.route("/api/icp/generate", methods=["POST"])
def icp_generate():
    """AI 生成 ICP 外包需求文档"""
    data = request.get_json(silent=True) or {}
    required = ["company_name", "site_name", "domain"]
    for field in required:
        if not data.get(field, "").strip():
            return jsonify({"ok": False, "error": f"缺少必填字段: {field}"})

    doc, err = _generate_icp_doc(data)
    if err:
        return jsonify({"ok": False, "error": err})
    return jsonify({"ok": True, "document": doc})


@app.route("/api/icp/feishu", methods=["POST"])
def icp_send_feishu():
    """把 ICP 需求文档推送到飞书群"""
    data = request.get_json(silent=True) or {}
    chat_id = data.get("chat_id", "").strip()
    if not chat_id:
        return jsonify({"ok": False, "error": "缺少 chat_id"})
    form_data = data.get("form_data", data)
    doc, err = _generate_icp_doc(form_data)
    if err:
        return jsonify({"ok": False, "error": err})

    # 飞书 post 格式:每段一个 paragraph
    lines = doc.split("\n")
    paragraphs = [[{"tag": "text", "text": line}] for line in lines if line.strip()]
    r = feishu_send_post(chat_id, "📄 ICP 备案外包需求文档", paragraphs)
    if not r.get("ok"):
        return jsonify({"ok": False, "error": r.get("error", "发送失败")})
    return jsonify({"ok": True, "sent": True})


# === 飞书 webhook ===
_PROCESSED_MSG_IDS = set()
_MAX_MSG_CACHE = 200


def _handle_feishu_message(text, chat_id):
    """处理飞书消息指令,异步调用(不阻塞 webhook 响应)"""
    text = (text or "").strip()
    try:
        if "帮助" in text or text.lower() in ("help", "?", "？"):
            feishu_send_text(chat_id,
                "🤖 法务审核助手 · 指令列表\n"
                "──────────────\n"
                "审核  — 查看审核示例报告\n"
                "模板  — 查看已入库合同模板\n"
                "帮助  — 显示本指令列表\n"
                "──────────────\n"
                "直接发送关键词即可,无需@")

        elif "审核" in text or "示例" in text:
            # 用 test_contract.txt 做示例审核
            import os
            test_file = os.path.join(os.path.dirname(__file__), "test_contract.txt")
            if os.path.exists(test_file):
                with open(test_file, "r", encoding="utf-8") as f:
                    contract_text = f.read()
                feishu_send_text(chat_id, "⏳ 正在审核示例合同,请稍候 30-60 秒...")
                result = _do_review(contract_text)
                if result.get("ok"):
                    stats = result.get("stats", {})
                    lines = [
                        f"📋 示例审核报告",
                        f"合同类型: {result.get('contract_type', '未知')}",
                        f"风险等级: {result.get('overall_risk', '未知')}",
                        f"统计: 通过 {stats.get('pass', 0)} / 警告 {stats.get('warn', 0)} / 不合规 {stats.get('fail', 0)} (共 {stats.get('total', 0)} 项)",
                        "",
                    ]
                    for item in result.get("items", [])[:5]:
                        status = item.get("status", "")
                        emoji = {"pass": "✅", "warn": "⚠️", "fail": "❌"}.get(status, "")
                        issue = item.get("issue", item.get("suggestion", ""))[:60]
                        lines.append(f"{emoji} {item.get('item', '')[:30]}: {issue}")
                    feishu_send_text(chat_id, "\n".join(lines))
                else:
                    feishu_send_text(chat_id, f"❌ 审核失败: {result.get('error', '未知错误')}")
            else:
                feishu_send_text(chat_id, "📄 未找到示例合同文件。")

        elif "模板" in text or "文档" in text:
            conn = sqlite3.connect(DB_PATH)
            try:
                c = conn.cursor()
                c.execute("SELECT doc_type, doc_name, chunk_count FROM documents ORDER BY id")
                rows = c.fetchall()
            finally:
                conn.close()
            if rows:
                lines = ["📚 已入库合同模板:"]
                for doc_type, doc_name, chunk_count in rows:
                    lines.append(f"  • {doc_name} ({doc_type}, {chunk_count} chunks)")
                feishu_send_text(chat_id, "\n".join(lines))
            else:
                feishu_send_text(chat_id, "📚 知识库暂无文档。")

        else:
            feishu_send_text(chat_id,
                f"收到: {text[:50]}\n发送\"帮助\"查看可用指令。")
    except Exception as e:
        try:
            feishu_send_text(chat_id, f"❌ 处理消息时出错: {e}")
        except Exception:
            pass


def _handle_feishu_file_message(msg_id, msg_type, content, chat_id, sender_open_id):
    """处理飞书 file/image 消息:下载 → 提取文本 → 分类+审核 → 回复

    msg_type: "file" 或 "image"
    content: 飞书消息 content(已解析的 dict)
    """
    import io as _io

    try:
        # Step 0: 提取 file_key / image_key
        if msg_type == "image":
            file_key = content.get("image_key", "")
            file_type = "image"
            filename = f"feishu_image_{file_key[:16]}.png"
        elif msg_type == "file":
            file_key = content.get("file_key", "")
            file_type = "file"
            filename = content.get("file_name", f"feishu_file_{file_key[:16]}")
        else:
            return

        if not file_key:
            feishu_send_text(chat_id, "❌ 无法获取文件 key")
            return

        # Step 1: 下载文件
        feishu_send_text(chat_id, "📥 正在下载文件...")
        dl = feishu_download_file(msg_id, file_key, file_type=file_type)
        if not dl.get("ok"):
            feishu_send_text(chat_id, f"❌ 文件下载失败: {dl.get('error', '未知错误')}")
            return
        file_bytes = dl["data"]

        # Step 2: 提取文本
        feishu_send_text(chat_id, f"📄 文件已下载({len(file_bytes)} 字节),文本提取中...")
        extracted = extract_text_from_upload(_io.BytesIO(file_bytes), filename)
        if isinstance(extracted, dict) and not extracted.get("ok", True):
            feishu_send_text(chat_id, f"❌ 文本提取失败: {extracted.get('error', '未知')}")
            return
        contract_text = extracted
        if len(contract_text.strip()) < 20:
            feishu_send_text(chat_id, "❌ 提取到的文本过短,无法审核(文件可能无文字或 OCR 失败)")
            return

        # Step 3: 审核
        feishu_send_text(chat_id, f"🤖 AI 审核中(约 1-3 分钟,合同长度 {len(contract_text)} 字)...")
        result = _do_review(contract_text)
        if not result.get("ok"):
            feishu_send_text(chat_id, f"❌ 审核失败: {result.get('error', '未知错误')}")
            return

        # Step 4: 持久化到 contracts 表
        contract_type = result.get("contract_type", "未知")
        overall_risk = result.get("overall_risk", "未知")
        stats = result.get("stats", {})
        try:
            conn = sqlite3.connect(DB_PATH)
            try:
                c = conn.cursor()
                c.execute("""CREATE TABLE IF NOT EXISTS contracts (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    filename TEXT,
                    contract_type TEXT,
                    overall_risk TEXT,
                    total_items INTEGER,
                    pass_count INTEGER,
                    warn_count INTEGER,
                    fail_count INTEGER,
                    contract_text TEXT,
                    review_result TEXT,
                    feishu_open_id TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )""")
                c.execute(
                    "INSERT INTO contracts (filename, contract_type, overall_risk, total_items, pass_count, warn_count, fail_count, contract_text, review_result, feishu_open_id) VALUES (?,?,?,?,?,?,?,?,?,?)",
                    (
                        filename,
                        contract_type,
                        overall_risk,
                        stats.get("total", 0),
                        stats.get("pass", 0),
                        stats.get("warn", 0),
                        stats.get("fail", 0),
                        contract_text[:5000],
                        json.dumps(result.get("items", []), ensure_ascii=False)[:5000],
                        sender_open_id,
                    ),
                )
                conn.commit()
            finally:
                conn.close()
        except Exception as e:
            print(f"[warn] contracts 入库失败: {e}")

        # Step 5: Bot 回复富文本
        risk_emoji = {"高": "🔴", "中": "🟡", "低": "🟢"}.get(overall_risk, "⚪")
        paragraphs = [
            [{"tag": "text", "text": f"📋 合同审核完成\n"}],
            [{"tag": "text", "text": f"文件: {filename[:50]}\n"}],
            [{"tag": "text", "text": f"合同类型: {contract_type}\n"}],
            [{"tag": "text", "text": f"风险等级: {risk_emoji} {overall_risk}\n"}],
            [{"tag": "text", "text": f"统计: ✅ 通过 {stats.get('pass',0)}  ⚠️ 警告 {stats.get('warn',0)}  ❌ 不合规 {stats.get('fail',0)} (共 {stats.get('total',0)} 项)\n"}],
            [{"tag": "text", "text": "─" * 20}],
        ]
        # 列出前 8 条审核结果
        for item in result.get("items", [])[:8]:
            status = item.get("status", "")
            emoji = {"pass": "✅", "warn": "⚠️", "fail": "❌"}.get(status, "•")
            item_name = item.get("item", "")[:25]
            issue = item.get("issue") or item.get("suggestion") or ""
            if issue:
                issue = issue[:50]
            paragraphs.append([{"tag": "text", "text": f"{emoji} {item_name}: {issue}\n"}])

        if len(result.get("items", [])) > 8:
            paragraphs.append([{"tag": "text", "text": f"...还有 {len(result['items'])-8} 项,查看网页版详情"}])

        paragraphs.append([{"tag": "text", "text": "\n发送\"帮助\"查看更多指令"}])
        feishu_send_post(chat_id, "📋 合同审核报告", paragraphs)

    except Exception as e:
        try:
            feishu_send_text(chat_id, f"❌ 处理文件时出错: {e}")
        except Exception:
            pass


@app.route("/webhook", methods=["POST"])
def webhook():
    """飞书事件订阅回调

    支持飞书 v2 事件格式(header.event_type)和 v1 格式。
    支持 Encrypt Key 加密模式。
    收到消息后立即返回 200,异步处理指令。
    """
    data = request.get_json(silent=True) or {}

    # 加密模式:飞书发 {"encrypt": "base64..."},需解密
    if "encrypt" in data and not data.get("challenge"):
        try:
            import base64, hashlib
            from Crypto.Cipher import AES
            encrypt_key = os.environ.get("LARK_ENCRYPT_KEY", "")
            if not encrypt_key:
                return jsonify({"error": "encrypt key not configured"}), 500
            key = hashlib.sha256(encrypt_key.encode("utf-8")).digest()
            enc = base64.b64decode(data["encrypt"])
            cipher = AES.new(key, AES.MODE_CBC, iv=enc[:16])
            decrypted = cipher.decrypt(enc[16:])
            # PKCS7 去填充
            pad = decrypted[-1]
            decrypted = decrypted[:-pad].decode("utf-8")
            data = json.loads(decrypted)
        except Exception as e:
            return jsonify({"error": f"decrypt failed: {e}"}), 500

    # challenge 校验(必须在 1 秒内返回)
    if "challenge" in data:
        return jsonify({"challenge": data["challenge"]})

    header = data.get("header", {})
    event = data.get("event", {})
    msg = event.get("message", {})
    if not msg:
        return jsonify({"ok": True})

    # 消息去重
    msg_id = msg.get("message_id", "")
    if msg_id and msg_id in _PROCESSED_MSG_IDS:
        return jsonify({"ok": True, "dedup": True})
    if msg_id:
        _PROCESSED_MSG_IDS.add(msg_id)
        if len(_PROCESSED_MSG_IDS) > _MAX_MSG_CACHE:
            _PROCESSED_MSG_IDS.pop()

    chat_id = msg.get("chat_id", "")
    content_str = msg.get("content", "{}")
    try:
        content = json.loads(content_str) if isinstance(content_str, str) else content_str
    except (json.JSONDecodeError, TypeError):
        content = {}
    text = content.get("text", "")
    msg_type = msg.get("message_type", "")
    sender = event.get("sender", {}).get("sender_id", {}).get("open_id", "")

    # 异步处理
    if chat_id:
        import threading
        if msg_type in ("image", "file"):
            t = threading.Thread(
                target=_handle_feishu_file_message,
                args=(msg_id, msg_type, content, chat_id, sender),
                daemon=True,
            )
            t.start()
        elif text:
            t = threading.Thread(target=_handle_feishu_message, args=(text, chat_id), daemon=True)
            t.start()

    return jsonify({"ok": True})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5003)), debug=os.environ.get("FLASK_DEBUG") == "1")
